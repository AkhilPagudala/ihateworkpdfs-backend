import uuid
import os
from fastapi import FastAPI, UploadFile, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware

import pdfplumber
from docx import Document
from docx.shared import Pt

app = FastAPI(title="ihateworkpdfs Backend - High Quality PDF to Word")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.post("/convert")
async def convert_pdf_to_word(file: UploadFile):

    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed.")

    file_id = str(uuid.uuid4())
    input_pdf = f"/tmp/{file_id}.pdf"

    original_name = os.path.splitext(file.filename)[0]
    output_docx = f"/tmp/{original_name}.docx"

    # Save uploaded file
    with open(input_pdf, "wb") as f:
        f.write(await file.read())

    # Create Word document
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    try:
        with pdfplumber.open(input_pdf) as pdf:
            for page in pdf.pages:
                lines = page.extract_text()

                if lines:
                    for line in lines.split("\n"):
                        doc.add_paragraph(line)
                else:
                    doc.add_paragraph(" ")

        doc.save(output_docx)

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"PDFPlumber conversion failed: {str(e)}"
        )

    if not os.path.exists(output_docx):
        raise HTTPException(status_code=500, detail="DOCX file not generated.")

    return FileResponse(
        output_docx,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{original_name}.docx"
    )
